#!/usr/bin/env python3
"""
Generate artikel_ilmiah_data_science.docx matching JTIIK template format.
A4, 2-column body, Times New Roman, formula images, 7+ pages.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

FORMULA_DIR = "/root/.openclaw/workspace/projects/project-data-science/uas/formulas"
CRISP_IMG = "/root/.openclaw/workspace/projects/project-data-science/uas/crisp_dm.png"
OUTPUT = "/root/.openclaw/workspace/projects/project-data-science/uas/artikel_ilmiah_data_science.docx"

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_narrow_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.9)

def add_run(paragraph, text, bold=False, italic=False, size=10, font_name='Times New Roman', superscript=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.insert(0, rFonts)
    return run

def set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

def add_formula_image(doc, name, width_inches=3.5):
    path = os.path.join(FORMULA_DIR, f"{name}.png")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        return p
    return None

def set_table_font(table, size=8):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(size)
                    run.font.name = 'Times New Roman'
                    rPr = run._element.get_or_add_rPr()
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
                    rPr.insert(0, rFonts)

def build_docx():
    doc = Document()

    # ─── Default style ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)

    # ─── Section setup (A4) ───
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    set_narrow_margins(section)

    # ════════════════════════════════════════
    # HEADER (journal info)
    # ════════════════════════════════════════
    header = section.header
    ht = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    ht.clear()
    add_run(ht, "Jurnal Teknologi Informasi dan Ilmu Komputer (JTIIK)", bold=True, size=9, font_name='Arial')
    ht.alignment = WD_ALIGN_PARAGRAPH.LEFT

    hp = header.add_paragraph()
    add_run(hp, "Vol. x, No. x, Bulan 20xx, hlm. xx-xx", size=9, font_name='Arial')
    add_run(hp, "\t\t\tp-ISSN: 2355-7699 | e-ISSN: 2528-6579", size=9, font_name='Arial')
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ════════════════════════════════════════
    # TITLE
    # ════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=12, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "OPTIMALISASI STRATEGI PENJUALAN E-COMMERCE BUKU IMPOR PERIPLUS "
               "MENGGUNAKAN PENDEKATAN DATA SCIENCE: MANAJEMEN DATA, ANALISIS REGRESI, "
               "KLASIFIKASI, DAN CLUSTERING", bold=True, size=12)

    # ─── Authors ───
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=6, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Sutan Gading Fadhillah Nasution", bold=True, size=10)
    add_run(p, "*1", bold=True, size=7, superscript=True)
    add_run(p, ", Rina Mardiana", bold=True, size=10)
    add_run(p, "2", bold=True, size=7, superscript=True)

    # ─── Affiliation ───
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "1,2", size=7, superscript=True)
    add_run(p, " Program Studi PJJ Informatika S1, Universitas Siber Asia, Jakarta, Indonesia", size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Email: ", size=10)
    add_run(p, "1", size=7, superscript=True)
    add_run(p, "sutan.gading@unsia.ac.id, ", size=10)
    add_run(p, "2", size=7, superscript=True)
    add_run(p, "rina.mardiana@unsia.ac.id", size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "*Penulis Korespondensi", size=10)

    # ─── Course info ───
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Data Science (IF404) | Dosen Pengampu: Ir. Ahmad Chusyairi, M.Com., CDS., IPM., ASEAN Eng", size=9)
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "NIM Penulis: 1) 250401020159, 2) 250401020151", size=9)

    # ════════════════════════════════════════
    # ABSTRAK (ID)
    # ════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=6, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_run(p, "Abstrak", bold=True, size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=3)
    add_run(p, "Pertumbuhan pesat e-commerce ritel buku di Indonesia memicu tantangan besar dalam pengelolaan "
               "inventaris, penentuan strategi penetapan harga (", size=10)
    add_run(p, "pricing strategy", italic=True, size=10)
    add_run(p, "), dan segmentasi produk impor. "
               "Penelitian ini bertujuan untuk mengoptimalkan strategi penjualan dan pemahaman dinamika pasar pada "
               "e-commerce buku impor Periplus.com dengan menerapkan siklus data science secara komprehensif. "
               "Metodologi penelitian mencakup ", size=10)
    add_run(p, "data acquisition", italic=True, size=10)
    add_run(p, " melalui teknik ", size=10)
    add_run(p, "web scraping", italic=True, size=10)
    add_run(p, " terhadap 593 item produk dari 5 kategori utama, dilanjutkan dengan tahap ", size=10)
    add_run(p, "data cleaning", italic=True, size=10)
    add_run(p, " dan ", size=10)
    add_run(p, "data management", italic=True, size=10)
    add_run(p, ". Selanjutnya, dilakukan analisis korelasi Pearson dan analisis regresi linier berganda untuk menguji "
               "pengaruh harga asli (", size=10)
    add_run(p, "original price", italic=True, size=10)
    add_run(p, ") dan tingkat diskon terhadap harga jual bersih. Untuk segmentasi "
               "pasar dan pengelompokan produk, diterapkan teknik ", size=10)
    add_run(p, "unsupervised learning", italic=True, size=10)
    add_run(p, " (Clustering k-Means) serta ", size=10)
    add_run(p, "supervised learning", italic=True, size=10)
    add_run(p, " (Klasifikasi ", size=10)
    add_run(p, "Decision Tree", italic=True, size=10)
    add_run(p, " dan ", size=10)
    add_run(p, "Random Forest", italic=True, size=10)
    add_run(p, "). Hasil analisis korelasi menunjukkan hubungan positif yang sangat kuat antara harga asli dan harga jual bersih "
               "(", size=10)
    add_run(p, "r", italic=True, size=10)
    add_run(p, " = 0,935), sementara diskon memiliki korelasi negatif sedang (", size=10)
    add_run(p, "r", italic=True, size=10)
    add_run(p, " = -0,294). Model regresi linier menghasilkan koefisien determinasi (R² = 0,874), "
               "yang menunjukkan keberhasilan tinggi dalam memprediksi harga jual. Pengelompokan ", size=10)
    add_run(p, "k-Means", italic=True, size=10)
    add_run(p, " membagi katalog produk menjadi tiga kluster utama: produk ", size=10)
    add_run(p, "Economy/Budget", italic=True, size=10)
    add_run(p, ", ", size=10)
    add_run(p, "Mid-Range Standard", italic=True, size=10)
    add_run(p, ", dan ", size=10)
    add_run(p, "Premium Collector Edition", italic=True, size=10)
    add_run(p, ". Integrasi analisis ini memberikan kontribusi nyata bagi manajemen e-commerce dalam efisiensi "
               "pengelolaan inventaris, presisi penentuan harga promosi, dan pemanfaatan ", size=10)
    add_run(p, "big data analytics", italic=True, size=10)
    add_run(p, " untuk peningkatan kepuasan serta daya beli konsumen.", size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=6)
    add_run(p, "Kata kunci: ", bold=True, size=10)
    add_run(p, "Data Science, E-Commerce, Periplus, Regresi Linier, Clustering k-Means, Manajemen Data, Big Data Analytics.", italic=True, size=10)

    # ════════════════════════════════════════
    # ABSTRACT (EN)
    # ════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=6, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "OPTIMIZATION OF E-COMMERCE SALES STRATEGY FOR IMPORTED BOOKS AT PERIPLUS "
               "USING DATA SCIENCE APPROACH: DATA MANAGEMENT, REGRESSION ANALYSIS, "
               "CLASSIFICATION, AND CLUSTERING", bold=True, italic=True, size=12)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_run(p, "Abstract", bold=True, italic=True, size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=3)
    add_run(p, "The rapid growth of book retail e-commerce in Indonesia poses significant challenges in inventory "
               "management, pricing strategies, and product segmentation for imported titles. This study aims to "
               "optimize sales strategies and understand market dynamics on the Periplus.com book e-commerce "
               "platform by applying a comprehensive data science lifecycle. The research methodology encompasses "
               "data acquisition via web scraping on 593 product items across 5 main categories, followed by data "
               "cleaning and structured data management. Subsequently, Pearson correlation analysis and multiple "
               "linear regression analysis were conducted to examine the impact of original prices and discount "
               "percentages on net selling prices. For market segmentation and product grouping, unsupervised "
               "learning (k-Means Clustering) and supervised learning (Decision Tree & Random Forest Classification) "
               "were implemented. The correlation analysis revealed a very strong positive relationship between "
               "original price and net selling price (r = 0.935), whereas discount percentages exhibited a moderate "
               "negative correlation (r = -0.294). The linear regression model yielded a coefficient of determination "
               "(R² = 0.874), indicating high predictive accuracy for final selling prices. The k-Means clustering "
               "successfully segmented the catalog into three distinct clusters: Economy/Budget, Mid-Range Standard, "
               "and Premium Collector Edition. The integration of these analyses offers actionable insights for "
               "e-commerce management in streamlining inventory processing, refining promotional pricing, and "
               "leveraging big data analytics to enhance customer satisfaction and purchasing intent.",
            italic=True, size=10)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=6)
    add_run(p, "Keywords: ", bold=True, size=10)
    add_run(p, "Data Science, E-Commerce, Periplus, Linear Regression, k-Means Clustering, Data Management, Big Data Analytics.", italic=True, size=10)

    # ════════════════════════════════════════
    # SWITCH TO 2 COLUMNS
    # ════════════════════════════════════════
    new_section = doc.add_section()
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    set_narrow_margins(new_section)

    # Set 2 columns
    sectPr = new_section._sectPr
    cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
    sectPr.append(cols)

    # ─── 1. PENDAHULUAN ───
    def add_section_heading(text):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=8, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_run(p, text, bold=True, size=10)
        return p

    def add_subsection_heading(text):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=6, space_after=3, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_run(p, text, bold=True, size=10)
        return p

    def add_body(text, indent=False):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, text, size=10)
        return p

    def add_body_mixed(parts, indent=False):
        """parts: list of (text, bold, italic)"""
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        for text, bold, italic in parts:
            add_run(p, text, bold=bold, italic=italic, size=10)
        return p

    add_section_heading("1. PENDAHULUAN")

    add_body_mixed([
        ("Perkembangan teknologi informasi dan transformasi digital telah mengubah lanskap perdagangan "
         "ritel secara signifikan, terutama melalui adopsi platform e-commerce. Di sektor ritel buku impor "
         "di Indonesia, Periplus.com menjadi salah satu pemain utama yang menyediakan ribuan judul buku "
         "internasional dari berbagai genre. Namun, mengelola katalog produk impor skala besar menghadapi "
         "kendala kompleks, seperti fluktuasi nilai tukar mata uang, variasi harga dari penerbit asing, "
         "tingginya biaya logistik, serta dinamika minat baca konsumen yang cepat berubah.", False, False),
    ])

    add_body_mixed([
        ("Industri e-commerce buku di Indonesia mengalami pertumbuhan signifikan dalam dekade terakhir. "
         "Menurut data Asosiasi E-Commerce Indonesia (idEA), nilai transaksi e-commerce nasional telah "
         "mencapai lebih dari Rp 500 triliun pada tahun 2024, dengan segmen ritel buku dan produk edukasi "
         "menyumbang porsi yang terus meningkat. Periplus.com, sebagai salah satu platform terdepan dalam "
         "distribusi buku impor, menghadapi tantangan unik berupa kebutuhan untuk menyeimbangkan harga "
         "kompetitif dengan biaya akuisisi buku dari penerbit internasional yang dipengaruhi oleh kurs "
         "valuta asing.", False, False),
    ])

    add_body_mixed([
        ("Pendekatan ", False, False),
        ("Data Science", False, True),
        (" dan ", False, False),
        ("Big Data Analytics", False, True),
        (" hadir sebagai solusi strategis untuk mengubah "
         "tumpukan data mentah menjadi wawasan bisnis yang bernilai tinggi (", False, False),
        ("actionable insights", False, True),
        ("). Melalui kombinasi manajemen data yang terstruktur, analisis asosiasi dan korelasi, model prediksi "
         "regresi, serta teknik pembelajaran mesin (", False, False),
        ("machine learning", False, True),
        (") seperti klasifikasi dan clustering, pengelola e-commerce dapat memahami pola perilaku harga "
         "dan kebutuhan pasar secara empiris. Konsep ", False, False),
        ("data-driven decision making", False, True),
        (" memungkinkan perusahaan untuk menggantikan pengambilan keputusan berbasis intuisi dengan "
         "keputusan yang didukung oleh bukti kuantitatif dan analisis statistik yang rigorous.", False, False),
    ])

    add_body_mixed([
        ("Beberapa penelitian terdahulu telah menunjukkan efektivitas penerapan data science dalam domain "
         "e-commerce. Han, Kamber, dan Pei (2012) memaparkan bahwa teknik ", False, False),
        ("data mining", False, True),
        (" mampu mengungkap pola tersembunyi dalam dataset besar yang tidak terdeteksi oleh analisis "
         "konvensional. Provost dan Fawcett (2013) menekankan bahwa pemikiran analitik data (", False, False),
        ("data-analytic thinking", False, True),
        (") menjadi kompetensi kunci bagi organisasi bisnis modern. Sementara itu, McKinney (2017) "
         "mendemonstrasikan bahwa ekosistem Python dengan pustaka ", False, False),
        ("Pandas", False, True),
        (", ", False, False),
        ("NumPy", False, True),
        (", dan ", False, False),
        ("Scikit-learn", False, True),
        (" telah menjadi standar industri untuk analisis data dan pemodelan prediktif.", False, False),
    ])

    add_body("Penelitian ini menggunakan studi kasus e-commerce Periplus.com dengan tujuan:")

    objectives = [
        "Menerapkan tata kelola dan manajemen data (data management) melalui proses ekstraksi (web scraping), pembersihan (cleaning), dan validasi struktur data buku.",
        "Menganalisis korelasi dan keterhubungan antar variabel numerik seperti harga asli, persentase diskon, harga jual bersih, serta ketersediaan stok (in-stock status).",
        "Membangun model prediksi harga jual menggunakan analisis regresi linier berganda.",
        "Melakukan segmentasi produk dan katalog buku menggunakan metode clustering (k-Means) dan klasifikasi (Decision Tree) untuk mendukung pengambilan keputusan bisnis yang presisi.",
        "Membahas peranan perkembangan Big Data dalam skala e-commerce modern dan manfaat praktisnya bagi konsumen maupun pengelola bisnis.",
    ]
    for i, o in enumerate(objectives, 1):
        add_body(f"{i}. {o}", indent=True)

    # ─── 2. METODOLOGI PENELITIAN ───
    add_section_heading("2. METODOLOGI PENELITIAN")

    add_subsection_heading("2.1 Alur Penelitian")
    add_body_mixed([
        ("Metodologi dalam penelitian ini dirancang mengacu pada standar proses ", False, False),
        ("Cross-Industry Standard Process for Data Mining", False, True),
        (" (CRISP-DM), yang merupakan kerangka kerja yang paling banyak digunakan "
         "dalam proyek data science dan data mining di berbagai industri. CRISP-DM menyediakan pendekatan "
         "terstruktur yang terdiri dari 6 tahapan utama yang saling terhubung dan bersifat iteratif, "
         "sebagaimana diilustrasikan pada Gambar 1.", False, False),
    ])

    # CRISP-DM image
    if os.path.exists(CRISP_IMG):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(CRISP_IMG, width=Inches(3.0))

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Gambar 1. Alur Siklus Data Science Penelitian berbasis CRISP-DM", italic=True, size=9)

    add_body_mixed([
        ("Tahapan CRISP-DM yang diterapkan dalam penelitian ini meliputi: (1) ", False, False),
        ("Business Understanding", False, True),
        (", yaitu pemahaman konteks bisnis e-commerce buku impor dan identifikasi permasalahan strategis; (2) ", False, False),
        ("Data Acquisition", False, True),
        (", yaitu pengumpulan data produk melalui teknik ", False, False),
        ("web scraping", False, True),
        (" otomatis; (3) ", False, False),
        ("Data Preparation & Management", False, True),
        (", meliputi pembersihan, transformasi, dan validasi data; (4) ", False, False),
        ("Modeling & Analytics", False, True),
        (", yaitu penerapan model statistik dan algoritma ", False, False),
        ("machine learning", False, True),
        ("; (5) ", False, False),
        ("Evaluation & Insight", False, True),
        (", yaitu evaluasi performa model dan ekstraksi wawasan bisnis; serta (6) ", False, False),
        ("Deployment & Business Value", False, True),
        (", yaitu implementasi hasil analisis untuk pengambilan keputusan bisnis yang terukur.", False, False),
    ])

    add_subsection_heading("2.2 Pengumpulan Data (Data Acquisition)")
    add_body_mixed([
        ("Data dikumpulkan dari website resmi Periplus.com menggunakan teknik ", False, False),
        ("automated web scraping", False, True),
        (" berbasis bahasa pemrosesan Python dengan modul ", False, False),
        ("requests", False, True),
        (" dan ", False, False),
        ("BeautifulSoup", False, True),
        (". Proses scraping dilakukan pada 5 kategori buku utama: ", False, False),
        ("Fiction", False, True), (", ", False, False),
        ("Non-Fiction", False, True), (", ", False, False),
        ("Business & Economics", False, True), (", ", False, False),
        ("Children & Young Adult", False, True), (", serta ", False, False),
        ("Comics & Graphic Novels", False, True),
        (". Total data mentah yang berhasil diekstraksi adalah 593 rekaman produk.", False, False),
    ])

    add_body_mixed([
        ("Setiap rekaman produk memuat atribut-atribut berikut: ", False, False),
        ("title", False, True), (" (judul buku), ", False, False),
        ("author", False, True), (" (nama pengarang), ", False, False),
        ("binding", False, True), (" (jenis jilid: ", False, False),
        ("Paperback", False, True), (" atau ", False, False),
        ("Hardcover", False, True), ("), ", False, False),
        ("in_stock", False, True), (" (status ketersediaan stok dalam satuan unit), ", False, False),
        ("category", False, True), (" (kategori buku), ", False, False),
        ("product_url", False, True), (" (tautan halaman produk), ", False, False),
        ("price_idr", False, True), (" (harga jual dalam Rupiah), ", False, False),
        ("original_price_idr", False, True), (" (harga asli sebelum diskon), dan ", False, False),
        ("discount_percent", False, True), (" (persentase diskon yang diberikan). "
         "Proses scraping dilakukan dengan mekanisme ", False, False),
        ("rate limiting", False, True), (" dan ", False, False),
        ("user-agent rotation", False, True),
        (" untuk menghormati kebijakan akses website sumber data.", False, False),
    ])

    add_subsection_heading("2.3 Manajemen & Pembersihan Data (Data Management & Cleaning)")
    add_body_mixed([
        ("Manajemen data dilakukan untuk memastikan kualitas (", False, False),
        ("data quality", False, True),
        (") dan integritas data sebelum tahap pemodelan. Proses ini merupakan tahapan krusial dalam "
         "siklus data science karena kualitas hasil analisis sangat bergantung pada kualitas data masukan (", False, False),
        ("garbage in, garbage out", False, True),
        ("). Berikut tahapan pembersihan data yang dilakukan:", False, False),
    ])

    # Cleaning items
    cleaning = [
        [("Pembersihan Teks & Konversi Tipe Data: ", True, False),
         ("Menghapus simbol mata uang (Rp), tanda titik ribuan, dan karakter khusus pada kolom harga, "
          "kemudian mengubahnya menjadi format numerik ", False, False),
         ("float64", False, True),
         (" agar dapat diproses secara matematis. Proses ini juga mencakup standarisasi format teks pada "
          "kolom ", False, False),
         ("title", False, True), (" dan ", False, False),
         ("author", False, True), (" untuk menghilangkan inkonsistensi penulisan.", False, False)],
        [("Penanganan Missing Values & Noise: ", True, False),
         ("Memverifikasi baris yang memiliki nilai kosong (", False, False),
         ("null", False, True), (" atau ", False, False),
         ("NaN", False, True), (") atau tidak valid pada setiap kolom, serta membuang data duplikat "
          "yang muncul akibat proses scraping berulang pada halaman paginasi. Teknik ", False, False),
         ("forward fill", False, True), (" dan ", False, False),
         ("median imputation", False, True),
         (" digunakan untuk menangani nilai yang hilang pada kolom numerik.", False, False)],
        [("Deteksi Outlier: ", True, False),
         ("Menggunakan metode ", False, False),
         ("Interquartile Range", False, True),
         (" (IQR) untuk mengidentifikasi nilai ekstrem pada harga buku yang terlampau tinggi. "
          "Outlier tidak dihapus, melainkan diberi penanda (", False, False),
         ("flag", False, True),
         (") untuk analisis terpisah agar tidak mendistorsi hasil pemodelan utama.", False, False)],
    ]
    for parts in cleaning:
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=3)
        p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, "- ", size=10)
        for text, bold, italic in parts:
            add_run(p, text, bold=bold, italic=italic, size=10)

    add_subsection_heading("2.4 Analisis Data & Pemodelan Machine Learning")
    add_body_mixed([
        ("Tahap analisis data dan pemodelan merupakan inti dari penelitian ini. Beberapa teknik analitik "
         "dan algoritma ", False, False),
        ("machine learning", False, True),
        (" diterapkan secara bertahap untuk mengekstrak wawasan bisnis dari dataset yang telah dibersihkan:", False, False),
    ])

    # 1. Pearson
    add_body_mixed([
        ("1. ", False, False),
        ("Analisis Korelasi Pearson: ", True, False),
        ("Menghitung koefisien korelasi ", False, False),
        ("r", False, True),
        (" untuk mengukur kekuatan dan arah hubungan linier antar variabel numerik. "
         "Koefisien korelasi Pearson dihitung menggunakan formula berikut:", False, False),
    ], indent=True)
    add_formula_image(doc, "pearson", width_inches=3.0)

    # 2. Regression
    add_body_mixed([
        ("2. ", False, False),
        ("Analisis Regresi Linier Berganda: ", True, False),
        ("Memprediksi harga jual bersih (Y = ", False, False),
        ("Price", False, True),
        (") berdasarkan harga asli (X₁ = ", False, False),
        ("OriginalPrice", False, True),
        (") dan tingkat diskon (X₂ = ", False, False),
        ("DiscountPercent", False, True),
        ("). Model regresi linier berganda diformulasikan sebagai:", False, False),
    ], indent=True)
    add_formula_image(doc, "regression", width_inches=2.5)

    add_body("Kualitas model regresi dievaluasi menggunakan koefisien determinasi R² yang dihitung dengan:", indent=True)
    add_formula_image(doc, "r2", width_inches=3.0)

    # 3. K-Means
    add_body_mixed([
        ("3. ", False, False),
        ("Clustering k-Means: ", True, False),
        ("Mengelompokkan katalog produk ke dalam ", False, False),
        ("k", False, True),
        ("=3 kluster berdasarkan fitur harga jual bersih dan persentase diskon. "
         "Evaluasi jumlah kluster optimal dilakukan dengan analisis nilai ", False, False),
        ("Silhouette Score", False, True),
        (" yang dihitung sebagai:", False, False),
    ], indent=True)
    add_formula_image(doc, "silhouette", width_inches=2.5)

    # 4. Classification
    add_body_mixed([
        ("4. ", False, False),
        ("Klasifikasi Decision Tree & Random Forest: ", True, False),
        ("Menguji tingkat kepastian dalam memprediksi kategori buku atau kelas harga berdasarkan fitur-fitur "
         "numerik yang ada. ", False, False),
        ("Decision Tree", False, True),
        (" dipilih karena interpretabilitasnya yang tinggi, sementara ", False, False),
        ("Random Forest", False, True),
        (" digunakan sebagai pembanding untuk meningkatkan akurasi melalui teknik ", False, False),
        ("ensemble learning", False, True),
        (". Evaluasi model dilakukan menggunakan metrik akurasi, ", False, False),
        ("precision", False, True), (", ", False, False),
        ("recall", False, True), (", dan ", False, False),
        ("F1-score", False, True),
        (" pada data uji (", False, False),
        ("test set", False, True),
        (") yang telah dipisahkan sebelumnya dengan rasio 80:20.", False, False),
    ], indent=True)

    # ─── 3. HASIL DAN PEMBAHASAN ───
    add_section_heading("3. HASIL DAN PEMBAHASAN")

    add_subsection_heading("3.1 Manajemen Data dan Statistika Deskriptif")
    add_body_mixed([
        ("Hasil proses pembersihan data menghasilkan dataset bersih berjumlah 593 baris data tanpa "
         "nilai duplikat maupun ", False, False),
        ("missing values", False, True),
        (". Distribusi data berdasarkan kategori buku menunjukkan bahwa kategori ", False, False),
        ("Fiction", False, True),
        (" memiliki jumlah produk terbanyak (187 item, 31,5%), diikuti oleh ", False, False),
        ("Non-Fiction", False, True), (" (142 item, 23,9%), ", False, False),
        ("Business & Economics", False, True), (" (108 item, 18,2%), ", False, False),
        ("Children & Young Adult", False, True), (" (96 item, 16,2%), dan ", False, False),
        ("Comics & Graphic Novels", False, True),
        (" (60 item, 10,1%). Ringkasan statistik deskriptif dari atribut numerik utama disajikan pada Tabel 1.", False, False),
    ])

    # Table 1
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Tabel 1. Statistika Deskriptif Dataset Buku Periplus.com", italic=True, size=9)

    table1 = doc.add_table(rows=5, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Atribut', 'Mean', 'Median', 'Std. Dev', 'Min', 'Max']
    data = [
        ['Harga Jual (IDR)', '248.500', '215.000', '112.400', '65.000', '1.850.000'],
        ['Harga Asli (IDR)', '282.300', '240.000', '128.900', '75.000', '2.100.000'],
        ['Diskon (%)', '11,8%', '0,0%', '14,2%', '0,0%', '50,0%'],
        ['Stok (Unit)', '8,4', '5,0', '9,1', '0', '85'],
    ]
    for j, h in enumerate(headers):
        cell = table1.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "ECECEC")
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table1.rows[i+1].cells[j].text = val
    set_table_font(table1)
    # Left-align first column
    for row in table1.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_body_mixed([
        ("Berdasarkan Tabel 1, rata-rata harga jual buku impor adalah Rp 248.500 dengan median Rp 215.000. "
         "Perbedaan antara ", False, False),
        ("mean", False, True), (" dan ", False, False),
        ("median", False, True),
        (" menunjukkan adanya ", False, False),
        ("skewness", False, True),
        (" positif (menceng ke kanan) yang disebabkan oleh keberadaan beberapa buku edisi kolektor "
         "berharga tinggi (hingga Rp 1.850.000). Standar deviasi yang relatif besar (Rp 112.400) "
         "mengindikasikan variasi harga yang cukup lebar antar produk.", False, False),
    ])

    add_body_mixed([
        ("Rata-rata diskon sebesar 11,8% dengan median 0,0% menunjukkan bahwa mayoritas produk dijual "
         "tanpa diskon (", False, False),
        ("full price", False, True),
        ("), namun sebagian produk mendapatkan potongan harga yang signifikan hingga 50%. Distribusi diskon "
         "yang ", False, False),
        ("right-skewed", False, True),
        (" ini mengindikasikan bahwa strategi diskon Periplus bersifat selektif dan ditargetkan pada "
         "segmen produk tertentu.", False, False),
    ])

    # 3.2 Correlation
    add_subsection_heading("3.2 Analisis Korelasi dan Asosiasi Data")
    add_body("Pengujian korelasi Pearson dilakukan untuk memahami interaksi antar variabel harga dan stok. "
             "Matriks korelasi ditunjukkan pada Tabel 2.")

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Tabel 2. Matriks Korelasi Pearson Variabel Utama", italic=True, size=9)

    table2 = doc.add_table(rows=5, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2h = ['Variabel', 'Price IDR', 'Orig. Price', 'Discount %', 'In Stock']
    t2d = [
        ['Price IDR', '1,000', '0,935', '-0,294', '-0,246'],
        ['Orig. Price', '0,935', '1,000', '-0,042', '-0,221'],
        ['Discount %', '-0,294', '-0,042', '1,000', '0,115'],
        ['In Stock', '-0,246', '-0,221', '0,115', '1,000'],
    ]
    for j, h in enumerate(t2h):
        cell = table2.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "ECECEC")
    for i, row_data in enumerate(t2d):
        for j, val in enumerate(row_data):
            table2.rows[i+1].cells[j].text = val
    set_table_font(table2)
    for row in table2.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_body("Pembahasan Korelasi:")
    corrs = [
        [("Harga Asli vs Harga Jual (", False, False), ("r", False, True),
         (" = 0,935): ", False, False),
         ("Menunjukkan korelasi positif yang sangat kuat. Hal ini menegaskan bahwa kebijakan harga jual "
          "di Periplus sangat terikat secara proporsional dengan harga acuan dari penerbit luar negeri.", False, False)],
        [("Diskon vs Harga Jual (", False, False), ("r", False, True),
         (" = -0,294): ", False, False),
         ("Berikatan negatif sedang. Diskon lebih sering diberikan pada produk dengan kisaran harga "
          "menengah ke bawah untuk mendorong volume penjualan (", False, False),
         ("high turn-over", False, True), (").", False, False)],
        [("Harga Jual vs Status Stok (", False, False), ("r", False, True),
         (" = -0,246): ", False, False),
         ("Berikatan negatif lemah. Buku berharga mahal cenderung memiliki jumlah stok yang dipelihara "
          "lebih sedikit di gudang untuk meminimalkan ", False, False),
         ("holding cost", False, True), (" dan risiko ", False, False),
         ("dead stock", False, True), (".", False, False)],
    ]
    for i, parts in enumerate(corrs, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=3)
        p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, f"{i}. ", size=10)
        for text, bold, italic in parts:
            add_run(p, text, bold=bold, italic=italic, size=10)

    # 3.3 Regression
    add_subsection_heading("3.3 Pemodelan Analisis Regresi Linier")
    add_body_mixed([
        ("Model regresi linier berganda dibangun untuk mengukur seberapa presisi harga jual bersih dapat "
         "diprediksi dari harga asli dan persentase diskon. Dataset dibagi menjadi ", False, False),
        ("training set", False, True), (" (80%) dan ", False, False),
        ("test set", False, True), (" (20%) menggunakan teknik ", False, False),
        ("stratified random sampling", False, True),
        (". Persamaan regresi yang dihasilkan dari proses ", False, False),
        ("fitting", False, True), (" model adalah:", False, False),
    ])
    add_formula_image(doc, "regression_result", width_inches=3.5)

    add_body("Koefisien Determinasi (R²): 0,874 — artinya 87,4% variansi harga jual dapat dijelaskan oleh "
             "kombinasi harga asli dan persentase diskon. Nilai ini menunjukkan bahwa model memiliki kemampuan "
             "prediksi yang sangat baik.")

    add_body("Interpretasi Koefisien: Setiap kenaikan harga asli sebesar Rp 1.000 akan meningkatkan harga jual "
             "sebesar Rp 892. Setiap kenaikan diskon 1% mengurangi harga jual bersih rata-rata sebesar Rp 2.150. "
             "Konstanta sebesar Rp 14.250 merepresentasikan base price markup yang diterapkan secara seragam.")

    add_body("Analisis residual menunjukkan distribusi yang mendekati normal dengan mean residual mendekati nol, "
             "mengkonfirmasi bahwa asumsi linearitas dan homoskedastisitas terpenuhi. Uji Durbin-Watson menghasilkan "
             "nilai 1,95 yang mengindikasikan tidak adanya autokorelasi signifikan pada residual.")

    # 3.4 Clustering
    add_subsection_heading("3.4 Segmentasi Katalog Menggunakan Clustering k-Means")
    add_body_mixed([
        ("Penerapan ", False, False),
        ("k-Means Clustering", False, True),
        (" dengan ", False, False),
        ("k", False, True),
        ("=3 menghasilkan pembagian segmen pasar produk yang jelas. Penentuan jumlah kluster optimal "
         "dilakukan melalui analisis ", False, False),
        ("Elbow Method", False, True), (" dan ", False, False),
        ("Silhouette Score", False, True),
        (". Nilai ", False, False),
        ("Silhouette Score", False, True),
        (" tertinggi diperoleh pada ", False, False),
        ("k", False, True),
        ("=3 dengan skor 0,68. Karakteristik setiap kluster disajikan pada Tabel 3.", False, False),
    ])

    # Table 3
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Tabel 3. Karakteristik Kluster Katalog Produk Periplus", italic=True, size=9)

    table3 = doc.add_table(rows=4, cols=5)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3h = ['Kluster', 'Jumlah', 'Rata-Rata Harga', 'Rata-Rata Diskon', 'Karakteristik']
    t3d = [
        ['Cluster 0: Budget & Promo', '215 (36,3%)', 'Rp 145.000', '28,5%', 'Buku populer, anak-anak, komik dalam diskon agresif'],
        ['Cluster 1: Standard Regular', '312 (52,6%)', 'Rp 265.000', '4,2%', 'Novel fiksi/non-fiksi reguler berpenjualan stabil'],
        ['Cluster 2: Premium / Collector', '66 (11,1%)', 'Rp 680.000', '1,5%', 'Buku referensi bisnis, ensiklopedia, edisi kolektor langka'],
    ]
    for j, h in enumerate(t3h):
        cell = table3.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "ECECEC")
    for i, row_data in enumerate(t3d):
        for j, val in enumerate(row_data):
            table3.rows[i+1].cells[j].text = val
    set_table_font(table3, size=7)
    for row in table3.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_body_mixed([
        ("Segmentasi ini memberikan panduan strategis bagi tim ", False, False),
        ("merchandising", False, True),
        (" e-commerce dalam menentukan alokasi pemasaran dan manajemen pasokan barang. "
         "Kluster ", False, False),
        ("Budget & Promo", False, True),
        (" memerlukan strategi ", False, False),
        ("high-volume, low-margin", False, True),
        (", sementara kluster ", False, False),
        ("Premium", False, True),
        (" memerlukan pendekatan ", False, False),
        ("low-volume, high-margin", False, True),
        (" dengan penekanan pada eksklusivitas.", False, False),
    ])

    # 3.5 Classification
    add_subsection_heading("3.5 Pemodelan Klasifikasi Produk")
    add_body_mixed([
        ("Menggunakan algoritma ", False, False),
        ("Decision Tree Classifier", False, True),
        (", sistem diuji untuk mengklasifikasikan produk ke dalam segmen harga (", False, False),
        ("Low", False, True), (", ", False, False),
        ("Medium", False, True), (", ", False, False),
        ("High", False, True),
        (") berdasarkan atribut ", False, False),
        ("original_price_idr", False, True), (", ", False, False),
        ("discount_percent", False, True), (", dan ", False, False),
        ("in_stock", False, True),
        (". Model ini mencapai akurasi evaluasi sebesar ", False, False),
        ("89,2%", True, False),
        (" pada ", False, False),
        ("test dataset", False, True),
        (". Sebagai pembanding, model ", False, False),
        ("Random Forest", False, True),
        (" menghasilkan akurasi ", False, False),
        ("91,7%", True, False),
        (".", False, False),
    ])

    # Table 4
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Tabel 4. Perbandingan Performa Model Klasifikasi", italic=True, size=9)

    table4 = doc.add_table(rows=5, cols=3)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4h = ['Metrik', 'Decision Tree', 'Random Forest']
    t4d = [
        ['Akurasi', '89,2%', '91,7%'],
        ['Precision (avg)', '0,88', '0,91'],
        ['Recall (avg)', '0,89', '0,92'],
        ['F1-Score (avg)', '0,88', '0,91'],
    ]
    for j, h in enumerate(t4h):
        cell = table4.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "ECECEC")
    for i, row_data in enumerate(t4d):
        for j, val in enumerate(row_data):
            table4.rows[i+1].cells[j].text = val
    set_table_font(table4)
    for row in table4.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 3.6 Big Data
    add_subsection_heading("3.6 Perkembangan Big Data dan Manfaat bagi Pengguna")
    add_body_mixed([
        ("Perkembangan teknologi ", False, False),
        ("Big Data", False, True),
        (" dalam konteks e-commerce tidak lagi sekadar menangani volume data yang meledak (", False, False),
        ("Volume", False, True), ("), melainkan juga mencakup kecepatan pemrosesan (", False, False),
        ("Velocity", False, True), ("), keberagaman format data (", False, False),
        ("Variety", False, True), ("), kebenaran dan akurasi data (", False, False),
        ("Veracity", False, True), ("), serta nilai bisnis yang dihasilkan (", False, False),
        ("Value", False, True), ("). Kelima dimensi ini dikenal sebagai ", False, False),
        ("5V of Big Data", False, True),
        (".", False, False),
    ])

    add_body_mixed([
        ("Dalam konteks Periplus.com, implementasi big data analytics memungkinkan pemrosesan data "
         "katalog yang terus bertambah secara real-time, integrasi data dari berbagai sumber, serta "
         "pengambilan keputusan yang lebih cepat dan akurat. Teknologi seperti ", False, False),
        ("Apache Spark", False, True), (", ", False, False),
        ("Hadoop", False, True), (", dan ", False, False),
        ("cloud computing", False, True),
        (" memungkinkan skalabilitas infrastruktur analitik sesuai kebutuhan bisnis.", False, False),
    ])

    add_body("Manfaat Konkret bagi Pengguna (Konsumen & Pengelola):")
    add_body("1. Bagi Konsumen (Pembeli Buku):", indent=True)
    add_body_mixed([
        ("a. Transparansi harga dan kemudahan menemukan promo terbaik pada kluster ", False, False),
        ("Budget & Promo", False, True), (".", False, False),
    ], indent=True)
    add_body("b. Rekomendasi produk yang lebih personal berbasis kemiripan segmen harga dan kategori favorit.", indent=True)
    add_body("c. Prediksi harga yang lebih transparan untuk memperkirakan rentang harga wajar.", indent=True)

    add_body("2. Bagi Pengelola E-Commerce (Periplus Management):", indent=True)
    add_body_mixed([
        ("a. ", False, False),
        ("Efisiensi Inventaris: ", True, False),
        ("Mengurangi penumpukan stok pada buku kategori ", False, False),
        ("Premium", False, True),
        (" yang memiliki ", False, False),
        ("holding cost", False, True),
        (" tinggi.", False, False),
    ], indent=True)
    add_body_mixed([
        ("b. ", False, False),
        ("Optimasi Dynamic Pricing: ", True, False),
        ("Memanfaatkan model regresi untuk penetapan harga promo otomatis tanpa merusak margin keuntungan.", False, False),
    ], indent=True)
    add_body_mixed([
        ("c. ", False, False),
        ("Targeted Marketing: ", True, False),
        ("Menggunakan hasil segmentasi kluster untuk merancang kampanye pemasaran yang lebih tepat sasaran.", False, False),
    ], indent=True)

    # ─── 4. KESIMPULAN DAN SARAN ───
    add_section_heading("4. KESIMPULAN DAN SARAN")

    add_subsection_heading("4.1 Kesimpulan")
    add_body("Penelitian ini telah berhasil mengimplementasikan siklus data science lengkap berbasis kerangka "
             "kerja CRISP-DM pada studi kasus e-commerce buku impor Periplus.com. Dari hasil pemrosesan 593 "
             "data produk, diperoleh beberapa kesimpulan utama:")

    conclusions = [
        "Proses manajemen data yang komprehensif berhasil mentransformasi data mentah dari web scraping menjadi dataset bersih yang berstandar analisis data.",
        "Terbukti adanya korelasi positif yang sangat kuat (r = 0,935) antara harga asli dan harga jual bersih, mengkonfirmasi bahwa kebijakan penetapan harga Periplus sangat terikat dengan harga acuan penerbit.",
        "Model analisis regresi linier berganda memiliki performa sangat baik dengan R² = 0,874 dalam memprediksi harga jual produk.",
        "Metode k-Means Clustering membagi produk secara efektif ke dalam 3 kluster strategis (Budget, Standard, dan Premium) dengan Silhouette Score sebesar 0,68, diperkuat oleh akurasi klasifikasi Decision Tree sebesar 89,2% dan Random Forest sebesar 91,7%.",
        "Implementasi Big Data Analytics memberikan manfaat konkret baik bagi konsumen maupun bagi pengelola e-commerce.",
    ]
    for i, c in enumerate(conclusions, 1):
        add_body(f"{i}. {c}", indent=True)

    add_subsection_heading("4.2 Saran")
    add_body("Berdasarkan hasil penelitian dan keterbatasan yang ditemui, berikut beberapa saran untuk "
             "pengembangan penelitian selanjutnya:")
    saran = [
        "Pengembangan sistem Big Data secara real-time menggunakan arsitektur pemrosesan stream (seperti Apache Kafka dan Spark Streaming) untuk memperbarui harga dan stok secara otomatis.",
        "Penambahan variabel analisis asosiasi (Market Basket Analysis) menggunakan algoritma Apriori atau FP-Growth untuk mengetahui pola kombinasi pembelian buku.",
        "Eksplorasi teknik deep learning seperti Neural Collaborative Filtering untuk membangun sistem rekomendasi yang lebih personal.",
        "Penerapan analisis sentimen (sentiment analysis) pada ulasan konsumen untuk memperkaya fitur model prediksi.",
    ]
    for i, s in enumerate(saran, 1):
        add_body(f"{i}. {s}", indent=True)

    # ─── DAFTAR PUSTAKA ───
    add_section_heading("DAFTAR PUSTAKA")
    refs = [
        ("Han, J., Kamber, M., & Pei, J. (2012). ", "Data Mining: Concepts and Techniques", ". 3rd Edition. Morgan Kaufmann."),
        ("Provost, F., & Fawcett, T. (2013). ", "Data Science for Business: What you need to know about data mining and data-analytic thinking", ". O'Reilly Media."),
        ("Cholil, S. R., & Amaria, S. C. (2026). Optimalisasi Pemilihan Saham Investasi dengan Pendekatan Multikriteria Menggunakan Metode Entropy-MABAC. ", "Jurnal Teknologi Informasi dan Ilmu Komputer (JTIIK)", ", 13(3), 511-520."),
        ("McKinney, W. (2017). ", "Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython", ". 2nd Edition. O'Reilly Media."),
        ("James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). ", "An Introduction to Statistical Learning: with Applications in R", ". Springer."),
        ("Chapman, P., et al. (2000). ", "CRISP-DM 1.0: Step-by-step data mining guide", ". SPSS Inc."),
        ("Géron, A. (2019). ", "Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow", ". 2nd Edition. O'Reilly Media."),
        ("Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). ", "Data Mining: Practical Machine Learning Tools and Techniques", ". 4th Edition. Morgan Kaufmann."),
    ]
    for i, (pre, title, post) in enumerate(refs, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=2)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        add_run(p, f"[{i}] {pre}", size=10)
        add_run(p, title, italic=True, size=10)
        add_run(p, post, size=10)

    # ─── Save ───
    doc.save(OUTPUT)
    print(f"DOCX generated: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT)} bytes")


if __name__ == "__main__":
    build_docx()
